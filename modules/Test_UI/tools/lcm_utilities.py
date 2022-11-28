import os
import lcm
import math
import time
import dataclasses
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from datetime import datetime
from pyproj import Transformer
from tqdm.notebook import trange, tqdm
from modules.control.LingK.Test_UI.tools.read_case import ReadCase
import modules.common.data_struct.lcm.localization.gps_imu_info_t as gps_imu_info_t  # gps通道
import modules.common.data_struct.proto.localization.gps_imu_info_pb2 as gps_imu_info_pb2


class ReadLcmAll:
    def __init__(self):
        self.lcmlog_dict = {'filename': '',
                            'date': '',
                            'start': '',
                            'finish': '',
                            'max_time': '',
                            "filename_pata": "",  # 输入文件的绝对路径
                            }

        self.case_info_dict = {'filename': '',
                               'id': '',
                               'modular': '',
                               'name': '',
                               'data_name': [],
                               'time_diff': [],
                               'lookup_data': '',
                               'lookup_data_path': ''
                               }

        self.gps_info_dict = {"utime": [],
                              "UTM_X": [],
                              "UTM_Y": [],
                              "wey4_velocity": [],
                              "max_stime": '',
                              }

    def create_lcm_dataset(self, lcmlog_file_path):
        lcmlog = lcm.EventLog(lcmlog_file_path, mode="r")
        buffer = [(e.eventnum, e.channel, e.data, e.timestamp) for e in tqdm(lcmlog)]
        lcmlog.close()
        data = pd.DataFrame(
            buffer, columns=["event-number", "channel", "data", "timestamp"]
        )
        data.filename = lcmlog_file_path.split("/")[-1]
        return data

    def convert_localtime_to_date_str(self, localtime):
        return f"{localtime.tm_year:02d}-{localtime.tm_mon:02d}-{localtime.tm_mday:02d}"

    def convert_localtime_to_time_str(self, localtime):
        return f"{localtime.tm_hour:02d}:{localtime.tm_min:02d}:{localtime.tm_sec:02d}"

    def get_lcmlog_datetime(self, lcm_dataframe):
        t0_date = time.localtime(lcm_dataframe["timestamp"].values[0] / 1e6)
        tf_date = time.localtime(lcm_dataframe["timestamp"].values[-1] / 1e6)
        date = self.convert_localtime_to_date_str(t0_date)
        start = self.convert_localtime_to_time_str(t0_date)
        end = self.convert_localtime_to_time_str(tf_date)

        max_time = "%.2f" % (lcm_dataframe["timestamp"].values[-1] / 1e6 - lcm_dataframe["timestamp"].values[0] / 1e6)
        self.lcmlog_dict['date'] = date
        self.lcmlog_dict['start'] = start
        self.lcmlog_dict['finish'] = end
        self.lcmlog_dict['max_time'] = max_time

        return date, start, end

    @dataclasses.dataclass
    class LcmlogMetaData:
        filename: str = ""
        date: str = ""
        start_time: str = ""
        end_time: str = ""
        summary_message: str = ""

    def generate_lcmlog_statistics(self, lcm_dataframe):
        date, start, end = self.get_lcmlog_datetime(lcm_dataframe)

        self.lcmlog_dict['filename'] = lcm_dataframe.filename
        self.case_info_dict['filename'] = lcm_dataframe.filename

        # prepare header
        table_template = (
            "{0:30}|{1:11}|{2:11}|{3:18}|{4:18}|{5:18}|"  # {index: column_width}
        )
        summary_msg = f"Filename: \t {lcm_dataframe.filename} \n"
        summary_msg += f"Date: \t\t {date} \n"
        summary_msg += f"Start: \t\t {start} \n"
        summary_msg += f"Finish: \t {end} \n\n"
        summary_msg += table_template.format(
            "LCM_CHANNEL",
            "MSG COUNT ",
            "RATE [Hz] ",
            "AVG INTERVAL [ms] ",
            "MIN INTERVAL [ms] ",
            "MAX INTERVAL [ms] ",
        )
        summary_msg += "\n"

        for lcm_channel in lcm_dataframe["channel"].unique():
            message_count = lcm_dataframe["channel"].value_counts()[lcm_channel]

            data = lcm_dataframe[lcm_dataframe["channel"].isin([lcm_channel])][
                ["timestamp"]
            ]
            timestamps = data["timestamp"].values / 10 ** 6
            valid_data = len(timestamps) > 1

            time_intervals = np.diff(timestamps) if valid_data else 0
            intv_avg = (
                np.around(1e3 * np.mean(time_intervals), decimals=3) if valid_data else 0
            )
            intv_min = (
                np.around(1e3 * np.min(time_intervals), decimals=3) if valid_data else 0
            )
            intv_max = (
                np.around(1e3 * np.max(time_intervals), decimals=3) if valid_data else 0
            )
            rate = np.around(1.0 / np.mean(time_intervals), decimals=0) if valid_data else 0

            message_count = str(message_count)
            intv_avg = str(intv_avg)
            intv_min = str(intv_min)
            intv_max = str(intv_max)
            rate = str(rate)

            summary_msg += (
                    table_template.format(
                        lcm_channel, message_count, rate, intv_avg, intv_min, intv_max
                    )
                    + "\n"
            )

        out = self.LcmlogMetaData()
        out.filename = lcm_dataframe.filename
        out.date = date
        out.start_time = start
        out.end_time = end
        out.summary_message = summary_msg
        return out

    def unpack_packets(self, lcmlog_dataframe, channel):
        df_lcm_channel = lcmlog_dataframe[lcmlog_dataframe['channel'].isin([channel])][['data', 'timestamp']]
        packets, packets_timestamp = zip(*df_lcm_channel.values)
        n_packets = len(df_lcm_channel.values)
        log_starting_time = np.array(packets_timestamp)
        timestamp = (np.array(packets_timestamp) - log_starting_time[0]) * 1e-6
        # timestamp = np.array(packets_timestamp)
        return log_starting_time, timestamp, packets

    def get_lcm_dataset(self, filename):
        self.lcmlog_dict["filename_path"] = filename
        lcmlog_dataframe = self.create_lcm_dataset(filename)
        lcmlog_metadata = self.generate_lcmlog_statistics(lcmlog_dataframe)
        print(lcmlog_metadata.summary_message)
        self.aaa= lcmlog_metadata.summary_message
        return lcmlog_dataframe

    def get_case_info(self):
        c = ReadCase()
        c.read('/home/wbl/2CAR/data_analysis/modules/control/LingK/Test_UI/statics/case_all.csv')
        self.case_info_dict['id'] = self.lcmlog_dict['filename'].split('_')[-2]
        for i, d in enumerate(c.cases_list_dict['id']):
            if d == self.lcmlog_dict['filename'].split('_')[-2]:
                self.case_info_dict['modular'] = (c.cases_list_dict['modular'][i])
                self.case_info_dict['name'] = (c.cases_list_dict['name'][i])
                break
        return self.case_info_dict

    def lookup_data(self):
        self.get_case_info()
        # 筛选同样ID的数据
        gvt_path = os.path.join(os.path.abspath(os.path.join(os.getcwd(), "..")),
                                str(self.lcmlog_dict['date'] + '/GVT'))
        for (dirpath, dirnames, filenames) in os.walk(gvt_path):
            for fn in filenames:
                if fn.split('_')[-2] == self.case_info_dict['id']:
                    self.case_info_dict['data_name'].append(fn)
        vut_time = datetime(2022, 11, 8, int(self.lcmlog_dict['filename'].split('_')[-1].split('.')[-3]),
                            int(self.lcmlog_dict['filename'].split('_')[-1].split('.')[-2]),
                            int(self.lcmlog_dict['filename'].split('_')[-1].split('.')[-1]))
        # 计算时间差值
        for d in self.case_info_dict['data_name']:
            gvt_time = datetime(2022, 11, 8, int(d.split('_')[-1].split('.')[-3]), int(d.split('_')[-1].split('.')[-2]),
                                int(d.split('_')[-1].split('.')[-1]))
            self.case_info_dict['time_diff'].append(abs(gvt_time - vut_time).seconds)

        # 选择对应数据输出绝对路径
        try:
            if min(self.case_info_dict['time_diff']) < 300:
                self.case_info_dict['lookup_data'] = (self.case_info_dict['data_name'][
                    (self.case_info_dict['time_diff'].index(min(self.case_info_dict['time_diff'])))])
                self.case_info_dict['lookup_data_path'] = (os.path.join(gvt_path, self.case_info_dict['lookup_data']))
            else:
                self.case_info_dict['lookup_data'] = '没找到合适的数据（时间差大于5分钟）'
        # 没有捕获内容，需要添加
        except:
            print('空值---------------------------------------------------------------')

        self.read_gps()

    def read_gps(self):
        transformer = Transformer.from_crs("EPSG:4326", "EPSG:32650")
        # 读取测试车数据
        yaw = []
        utm_x = []
        utm_y = []
        lcmlog_dataframe = self.get_lcm_dataset(self.lcmlog_dict["filename_path"])  # 这里输入了两次  需要解决
        channel_starting_time, timestamp, packets = self.unpack_packets(lcmlog_dataframe, "abL10n")  # GPS通道数据
        gps = gps_imu_info_pb2.gpsImu()
        for packet in packets:
            gps.ParseFromString(packet)
            yaw.append(gps.yaw)
            x, y = transformer.transform(gps.latitude, gps.longitude)
            utm_x.append(x)
            utm_y.append(y)
        utime_4 = []
        utm_4x = []
        utm_4y = []
        stime = []
        max_utime_4 = []
        log = lcm.EventLog(self.case_info_dict['lookup_data_path'], "r")
        for event in log:
            if event.channel == "GPS_DATA":
                msg = gps_imu_info_t.decode(event.data)
                max_utime_4.append(msg.utime)
                if min(channel_starting_time) < msg.utime < max(channel_starting_time):
                    utime_4.append(msg.utime)
                    x, y = transformer.transform(msg.latitude, msg.longitude)
                    utm_4x.append(x)
                    utm_4y.append(y)
                    self.gps_info_dict['wey4_velocity'].append(msg.velocity * 3.6)

        for t in utime_4:
            stime.append((t - utime_4[0]) / 1000000)
        UTM_X = []
        UTM_Y = []
        try:
            for i in range(len(utm_4x)):
                aa = utm_4x[i] - utm_x[i]  # 配合车－实验车XY  构建试验车坐标原点
                bb = utm_4y[i] - utm_y[i]
                xq = aa * math.cos(-yaw[i]) + bb * math.sin(-yaw[i])
                yq = bb * math.cos(-yaw[i]) - aa * math.sin(-yaw[i])
                UTM_X.append(xq)
                UTM_Y.append(yq)

        except IndexError:
            UTM_X.clear()
            UTM_Y.clear()
            stime.clear()
            stime = timestamp
            for i in range(len(utm_x)):
                aa = utm_4x[i] - utm_x[i]
                bb = utm_4y[i] - utm_y[i]
                xq = aa * math.cos(-yaw[i]) + bb * math.sin(-yaw[i])
                yq = bb * math.cos(-yaw[i]) - aa * math.sin(-yaw[i])
                UTM_X.append(xq)
                UTM_Y.append(yq)
        self.gps_info_dict['utime'] = stime
        self.gps_info_dict['UTM_X'] = UTM_X
        self.gps_info_dict['UTM_Y'] = UTM_Y
        max_stime = []
        for t in max_utime_4:
            max_stime.append((t - max_utime_4[0]) / 1000000)
        self.gps_info_dict['max_stime'] = "%.2f" % max(max_stime)

    def word_docx(self, path1, path2):
        self.document = Document()
        self.document.add_heading(
            '睿蓝_%s_%s' % (self.case_info_dict['modular'].split('/')[1],
                          self.case_info_dict['modular'].split('/')[2]), 0)
        self.document.add_heading('1.测试场景数据简述', 1)
        self.document.add_paragraph('测试日期：%s' % (self.lcmlog_dict['date']))
        self.document.add_paragraph('测试开始时间：%s' % (self.lcmlog_dict['start']))
        self.document.add_paragraph('测试车辆：%s' % (self.lcmlog_dict['filename'].split('_')[1]))
        # self.document.add_paragraph('测试结束时间：%s' % (self.lcmlog_dict['finish']))    #结束时间
        self.document.add_paragraph('测试用例名称：%s' % (self.case_info_dict['name']))
        self.document.add_paragraph('测试用例模块：%s' % (self.case_info_dict['modular']))

        # 拉满列表
        self.table = self.document.add_table(rows=1, cols=4)
        self.hdr_cells = self.table.rows[0].cells
        self.hdr_cells[0].text = '数据名称'
        self.hdr_cells[1].text = '测试车辆'
        self.hdr_cells[2].text = '测试人员'
        self.hdr_cells[3].text = '数据时长'

        # 测试车
        self.row_cells = self.table.add_row().cells
        self.row_cells[0].text = self.lcmlog_dict['filename']
        self.row_cells[1].text = self.lcmlog_dict['filename'].split('_')[1]
        self.row_cells[2].text = self.lcmlog_dict['filename'].split('_')[0]
        self.row_cells[3].text = self.lcmlog_dict['max_time']
        # 配车车   没有考虑无配合车数据场景  -------------------------
        self.row_cells = self.table.add_row().cells
        self.row_cells[0].text = self.case_info_dict['lookup_data']
        self.row_cells[1].text = self.case_info_dict['lookup_data'].split('_')[1]
        self.row_cells[2].text = self.case_info_dict['lookup_data'].split('_')[0]
        self.row_cells[3].text = str(self.gps_info_dict['max_stime'])

        self.document.add_heading('2.测试场景', 1)
        self.document.add_picture(path1, width=Inches(6.0))
        self.document.add_heading('3.测试数据分析', 1)
        self.document.add_picture(path2, width=Inches(6.0))
        self.document.add_heading('4.测试结论', 1)
        self.document.add_paragraph('  经测试，测试车数据01_LingK1_DOW_776_7153_16.49.34,功能触发几次，行驶了多少km，是否符合')
        # self.document.add_page_break()
        self.document.save('%s.docx' % (self.case_info_dict['filename']))


if __name__ == '__main__':
    APP = ReadLcmAll()
    APP.get_lcm_dataset(
        '/home/wbl/2CAR/data_analysis/modules/control/LingK/Test_UI/2022-11-14/VUT/01_LiK2_DOW_457_7296_15.13.05')
    lookup = APP.lookup_data()
    print(APP.case_info_dict)
    APP.word_docx('/home/wbl/2CAR/data_analysis/LingK13.55.30_Fusion.jpg',
                  '/home/wbl/2CAR/data_analysis/LingK13.55.30_Fusion.jpg')
