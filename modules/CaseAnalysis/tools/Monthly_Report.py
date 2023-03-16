""""
拉取mysql数据，通过里面的信息产生月度报告，生成world文档。
拉取数据    选择规定日期内的数据
测试用例执行结结果 、通过个数、失败个数、通过率；失败的测试日期、数据名称、用例抬头
按用例模块统计  条目   值   百分比
按用例执行人统计   条目   值   百分比
按车辆统计    车辆    值  百分比
本月度组织测试场次，分布情况，所占百分比；
末尾详细描述   详细情况
"""
import mysqlpy
from docx import Document
import matplotlib.pyplot as plt
from docx.shared import Inches


class MonthlyReport:
    def __int__(self):
        pass

    def report_forms(self):
        self.reports = []  # 所有数据
        self.fail = []  # 失败的数据
        # 获取数据库信息
        report = mysqlpy.check_mysql()
        # 筛选规定日期数据 将数据存入列表中
        for i in report:
            if int(str(i[1]).split('-')[1]) == 12:
                self.reports.append(i)
        # 按用例结果统计  条目  值   百分比   (数据个数、失败用例)
        for i in self.reports:
            if i[-1] == '失败':
                self.fail.append(i)
        # 绘图
        labels = 'fail', 'pass'  # 定义标签
        sizes = [len(self.fail) / len(self.reports) * 100, 100 - (len(self.fail) / len(self.reports) * 100)]  # 每一块的比例
        colors = ['green', 'red']  # 每一块的颜色
        explode = (0, 0)  # 突出显示，这里仅仅突出显示第二块（即'Hogs'）
        plt.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', shadow=True, startangle=90)
        plt.axis('equal')  # 显示为圆（避免比例压缩为椭圆）
        plt.savefig('case.jpg')
        # 总结
        self.case_report = (
            '    测试月份是12月，总共执行{}次测试，通过{}次，失败{}次。'.format(len(self.reports),
                                                                        len(self.reports) - len(
                                                                            self.fail),
                                                                        len(self.fail)))

        # 按照测试车辆统计   个车测试时长   自动驾驶里程
        test_time = 0
        test_mileage = 0
        for i in self.reports:
            if i[3] == 'LiK2':
                test_time += i[9]
                test_mileage += i[10]
        # print(test_time, test_mileage)
        self.car_report = ('    测试车辆LiK2，数据时常{}分钟，测试里程{}m。'.format(test_time, test_mileage))

        # 按照人员划分
        hjy = 0
        jxx = 0
        wbl = 0
        for i in self.reports:
            if i[2] == '02':
                jxx += 1
            elif i[2] == '01':
                hjy += 1
            elif i[3] == '03':
                wbl += 1
        self.person = ('    测试人员分布情况，姜晓序执行{}次，胡靖宇执行{}次，王博伦执行{}次。'.format(jxx, hjy, wbl))

    def word_docx(self):
        self.report_forms()
        self.document = Document()
        self.document.add_heading('睿蓝项目_12月份 月度报告汇总', 0)
        self.document.add_heading('概述', 1)
        self.document.add_paragraph(self.case_report)
        self.document.add_paragraph(self.car_report)
        self.document.add_paragraph(self.person)
        self.document.add_heading('测试用例执行情况', 1)
        self.document.add_picture('/home/dell/APP/CaseAnalysis/tools/case.jpg', width=Inches(6.0))
        self.document.add_paragraph('\n测试用例详细执行情况如下表所示：\n')
        # 列表包含数据名称  抬头  通过或失败 三项
        self.table = self.document.add_table(rows=1, cols=3)
        self.hdr_cells = self.table.rows[0].cells
        self.hdr_cells[0].text = '数据名称'
        self.hdr_cells[1].text = '用例名称'
        self.hdr_cells[2].text = '测试结果'
        #绘制表格
        case_len = len(self.reports)
        while case_len > 0:
            self.row_cells = self.table.add_row().cells
            self.row_cells[0].text = str(self.reports[int(case_len) - 1][6])
            self.row_cells[1].text = str(self.reports[int(case_len) - 1][8])
            self.row_cells[2].text = str(self.reports[int(case_len) - 1][-1])
            case_len -= 1

        self.document.add_heading('车俩使用情况', 1)
        self.document.add_paragraph('\n测试用车详细情况如下表所示：\n')
        self.table = self.document.add_table(rows=1, cols=5)
        self.hdr_cells = self.table.rows[0].cells
        self.hdr_cells[0].text = '数据名称'
        self.hdr_cells[1].text = '用例名称'
        self.hdr_cells[2].text = '测试车辆'
        self.hdr_cells[3].text = '数据时常（s）'
        self.hdr_cells[4].text = '行驶里程(m)'
        # 绘制表格
        case_len = len(self.reports)
        while case_len > 0:
            self.row_cells = self.table.add_row().cells
            self.row_cells[0].text = str(self.reports[int(case_len) - 1][6])
            self.row_cells[1].text = str(self.reports[int(case_len) - 1][8])
            self.row_cells[2].text = str(self.reports[int(case_len) - 1][3])
            self.row_cells[3].text = str(self.reports[int(case_len) - 1][9])
            self.row_cells[4].text = str(self.reports[int(case_len) - 1][10])
            case_len -= 1
        self.document.add_heading('执行人情况', 1)
        self.document.add_paragraph('\n测试执行人详细执行情况如下表所示：\n')
        self.table = self.document.add_table(rows=1, cols=3)
        self.hdr_cells = self.table.rows[0].cells
        self.hdr_cells[0].text = '数据名称'
        self.hdr_cells[1].text = '用例名称'
        self.hdr_cells[2].text = '测试人员'
        # 绘制表格
        case_len = len(self.reports)
        while case_len > 0:
            self.row_cells = self.table.add_row().cells
            self.row_cells[0].text = str(self.reports[int(case_len) - 1][6])
            self.row_cells[1].text = str(self.reports[int(case_len) - 1][8])
            self.row_cells[2].text = str(self.reports[int(case_len) - 1][2])
            case_len -= 1
        self.document.save('hello_world.docx')


if __name__ == '__main__':
    app = MonthlyReport()
    app.word_docx()
